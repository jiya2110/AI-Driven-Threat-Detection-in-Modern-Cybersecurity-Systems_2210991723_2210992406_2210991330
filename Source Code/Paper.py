from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# =========================
# PAGE SETUP (IEEE STRICT)
# =========================
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21)

section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

# 2 COLUMN
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '2')
cols.set(qn('w:space'), '720')

# =========================
# FONT
# =========================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)

# =========================
# TITLE
# =========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
"Artificial Intelligence Driven Threat Detection in Modern Cybersecurity Systems Using Hybrid Machine Learning Approaches"
)
run.bold = True
run.font.size = Pt(24)

# =========================
# AUTHORS (IEEE STYLE BLOCK)
# =========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(
"Jashit Goyal\nDept. of Computer Science\n[Your University]\nIndia\nemail@example.com\n\n"
"Co-Author Name\nDept. of Computer Science\n[University]\nIndia\nemail@example.com"
)

# =========================
# ABSTRACT
# =========================
p = doc.add_paragraph()
run = p.add_run(
"Abstract—Modern cybersecurity systems face increasingly sophisticated attacks that cannot be effectively detected using traditional signature-based approaches. "
"This paper presents a comprehensive Artificial Intelligence-driven framework for threat detection utilizing machine learning and hybrid anomaly detection techniques. "
"The NSL-KDD dataset is used to evaluate multiple classification models including Random Forest, Support Vector Machine, Decision Tree, XGBoost, and Artificial Neural Networks. "
"Experimental results demonstrate that XGBoost achieves superior precision (96.78%) and strong overall accuracy, while Decision Tree performs best in multi-class classification. "
"However, moderate recall values indicate limitations in detecting all attack instances. "
"To address this issue, a hybrid model integrating Autoencoder-based reconstruction error and Isolation Forest anomaly detection is proposed. "
"The results indicate that combining supervised and unsupervised learning enhances detection capability, particularly for previously unseen threats, making the system more robust for real-world deployment."
)
run.bold = True

# KEYWORDS
doc.add_paragraph(
"Keywords—Cybersecurity, Intrusion Detection, Machine Learning, Hybrid AI, NSL-KDD"
)

# =========================
# SECTION FUNCTION
# =========================
def add_section(title, content):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    doc.add_paragraph(content)

# =========================
# I. INTRODUCTION
# =========================
add_section(
"I. INTRODUCTION",
"The rapid digitization of modern systems has significantly increased the vulnerability of networks to cyberattacks. "
"Traditional intrusion detection systems rely on predefined signatures, which limits their effectiveness against evolving threats. "
"Artificial Intelligence enables adaptive detection by learning patterns from data. "
"This study evaluates machine learning approaches and introduces a hybrid model to improve detection performance. "
"With increasing attack complexity, there is a need for intelligent systems capable of detecting both known and unknown threats."
)

# =========================
# II. LITERATURE REVIEW
# =========================
add_section(
"II. LITERATURE REVIEW",
"Previous research has demonstrated the effectiveness of machine learning techniques in intrusion detection systems. "
"However, several limitations exist. Most studies rely on outdated datasets such as NSL-KDD, which do not reflect modern traffic patterns. "
"Additionally, models often suffer from overfitting and poor generalization. "
"Class imbalance leads to poor detection of rare attacks such as U2R and R2L. "
"Furthermore, supervised models cannot detect zero-day attacks due to reliance on labeled data. "
"These limitations highlight the need for hybrid approaches."
)

# =========================
# III. METHODOLOGY
# =========================
add_section(
"III. METHODOLOGY",
"The proposed system evaluates multiple machine learning models including Random Forest, Support Vector Machine, Decision Tree, XGBoost, K-Nearest Neighbors, Naive Bayes, and Artificial Neural Networks. "
"Binary and multi-class classification tasks are performed. "
"Feature scaling and preprocessing are applied to improve model performance. "
"A hybrid anomaly detection model is introduced to enhance detection of unknown threats."
)

# FIGURE PLACEHOLDER (TOP/BOTTOM RULE)
fig = doc.add_paragraph()
fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig.add_run("Fig. 1. System Architecture of Proposed Model")

# =========================
# IV. RESULTS
# =========================
add_section(
"IV. RESULTS",
"Experimental results indicate that XGBoost achieves the highest precision (96.78%) with strong accuracy. "
"Decision Tree performs best in multi-class classification. "
"Naive Bayes shows poor performance due to simplifying assumptions. "
"The results highlight trade-offs between detection rate and false positive rate."
)

fig = doc.add_paragraph()
fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig.add_run("Fig. 2. Model Performance Comparison")

# =========================
# V. DISCUSSION
# =========================
add_section(
"V. DISCUSSION",
"The results indicate that while machine learning models achieve high precision, recall remains moderate. "
"This suggests that many attacks remain undetected. "
"The hybrid model improves detection by identifying anomalies beyond labeled data. "
"This is critical for detecting zero-day attacks."
)

# =========================
# VI. FUTURE WORK
# =========================
add_section(
"VI. FUTURE WORK",
"Future work includes integration of deep learning models such as LSTM and CNN for temporal analysis. "
"Use of modern datasets such as CIC-IDS2017 is recommended. "
"Explainable AI techniques can improve interpretability of decisions. "
"Real-time deployment and federated learning are promising directions."
)

# =========================
# VII. CONCLUSION
# =========================
add_section(
"VII. CONCLUSION",
"This study demonstrates that AI-driven approaches significantly improve cybersecurity threat detection. "
"While traditional models achieve strong precision, hybrid approaches enhance robustness and detection of unknown threats. "
"The proposed system provides a scalable and effective solution for modern cybersecurity challenges."
)

# =========================
# REFERENCES (STRICT IEEE)
# =========================
p = doc.add_paragraph()
run = p.add_run("REFERENCES")
run.bold = True

doc.add_paragraph("[1] M. Tavallaee et al., 'A detailed analysis of the KDD Cup 99 dataset,' IEEE, 2009.")
doc.add_paragraph("[2] I. Sharafaldin et al., 'Toward generating a new intrusion detection dataset,' ICISSP, 2018.")
doc.add_paragraph("[3] A. Buczak and E. Guven, 'A survey of ML methods for cybersecurity,' IEEE, 2019.")
doc.add_paragraph("[4] J. Saxe and K. Berlin, 'Deep neural network-based malware detection,' IEEE, 2017.")
doc.add_paragraph("[5] X. Chen et al., 'Adversarial malware attacks,' IEEE, 2018.")

# =========================
# SAVE
# =========================
doc.save("IEEE_CAMERA_READY.docx")

print("🔥 IEEE CAMERA-READY PAPER GENERATED SUCCESSFULLY!")